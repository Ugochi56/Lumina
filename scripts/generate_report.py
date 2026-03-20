import os
import psycopg2
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from dotenv import load_dotenv

# Load Environment Variables
load_dotenv('.env')

def download_image(url, filepath):
    try:
        response = requests.get(url, stream=True, timeout=15)
        if response.status_code == 200:
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(8192):
                    f.write(chunk)
            
            # Verify image integrity to prevent ReportLab crashes
            from PIL import Image
            try:
                img = Image.open(filepath)
                img.verify()
                img.close()
                
                # Full decode and force normalize to clean JPEG 
                img = Image.open(filepath)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                img.save(filepath, 'JPEG')
                img.close()
                return True
            except Exception as img_err:
                print(f"Image {filepath} is truncated/corrupted: {img_err}")
                if os.path.exists(filepath): os.remove(filepath)
                return False
                
    except Exception as e:
        print(f"Failed to download {url}: {e}")
    
    if os.path.exists(filepath): os.remove(filepath)
    return False

def generate_reports():
    print("Generating Advanced Evaluation Reports...")

    DATABASE_URL = os.environ.get('DATABASE_URL')
    if not DATABASE_URL:
        print("Database URL not found.")
        return

    try:
        conn = psycopg2.connect(DATABASE_URL)
        cur = conn.cursor()
        
        # Pull Data into Pandas DataFrame
        query = """
            SELECT *
            FROM photos 
            WHERE recommended_tool IS NOT NULL 
                AND ssim_score IS NOT NULL 
                AND brisque_score IS NOT NULL
                AND brisque_score > 0
        """
        cur.execute(query)
        rows = cur.fetchall()
        cols = [desc[0] for desc in cur.description]
        df = pd.DataFrame(rows, columns=cols)
        
        # Rename recommended_tool to tool
        if 'recommended_tool' in df.columns:
            df.rename(columns={'recommended_tool': 'tool'}, inplace=True)
        
        # Cast metrics to float as psycopg2 fetches NUMERIC as decimal.Decimal
        for col in ['ssim_score', 'brisque_score', 'processing_time_ms']:
            if col in df.columns:
                df[col] = df[col].astype(float)
        
        cur.close()
        conn.close()
        
        if df.empty:
            print("Not enough evaluation data with SSIM/BRISQUE scores to generate a report.")
            return

        # Ensure directory
        os.makedirs('reports', exist_ok=True)
        os.makedirs('temp_eval', exist_ok=True)

        print("Generating charts...")
        generate_charts(df)
        
        print("Downloading case study images...")
        case_studies = get_case_studies(df)
        download_case_study_images(case_studies)

        print("Generating PDF Report...")
        generate_pdf_report(df, case_studies)

        print("Generating Word Document Report...")
        generate_word_report(df, case_studies)
        
        print(f"Reports Successfully Generated:\\n - reports/lumina_eval_report.pdf\\n - reports/lumina_eval_report.docx")

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error generating report: {e}")

def generate_charts(df):
    plt.style.use('dark_background')
    colors = {'upscale': '#e85d04', 'restore': '#4361ee', 'edit': '#2a9d8f', 'lowlight': '#e9c46a'}
    
    # 1. Bar chart SSIM
    plt.figure(figsize=(8, 6), facecolor='#03071e')
    ax = plt.axes()
    ax.set_facecolor('#03071e')
    avg_ssim = df.groupby('tool')['ssim_score'].mean()
    color_map = [colors.get(t, '#888888') for t in avg_ssim.index]
    avg_ssim.plot(kind='bar', color=color_map, edgecolor='none')
    plt.title('Average SSIM Score per Tool Type')
    plt.ylabel('SSIM Score')
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig('temp_eval/plot_ssim.png', facecolor='#03071e')
    plt.close()

    # 2. Bar chart BRISQUE
    plt.figure(figsize=(8, 6), facecolor='#03071e')
    ax = plt.axes()
    ax.set_facecolor('#03071e')
    avg_brisque = df.groupby('tool')['brisque_score'].mean()
    color_map = [colors.get(t, '#888888') for t in avg_brisque.index]
    avg_brisque.plot(kind='bar', color=color_map, edgecolor='none')
    plt.title('Average BRISQUE Score per Tool Type\\n(Lower BRISQUE is better)')
    plt.ylabel('BRISQUE Score')
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig('temp_eval/plot_brisque.png', facecolor='#03071e')
    plt.close()

    # 3. Scatter plot
    plt.figure(figsize=(8, 6), facecolor='#03071e')
    ax = plt.axes()
    ax.set_facecolor('#03071e')
    df['processing_time_sec'] = df['processing_time_ms'] / 1000.0
    for tool in df['tool'].unique():
        subset = df[df['tool'] == tool]
        plt.scatter(subset['processing_time_sec'], subset['ssim_score'], label=tool, color=colors.get(tool, '#888888'), alpha=0.7)
    plt.title('SSIM Score vs Processing Time')
    plt.xlabel('Processing Time (Seconds)')
    plt.ylabel('SSIM Score')
    plt.legend()
    plt.tight_layout()
    plt.savefig('temp_eval/plot_scatter.png', facecolor='#03071e')
    plt.close()

    # 4. Latency Distribution
    plt.figure(figsize=(8, 6), facecolor='#03071e')
    ax = plt.axes()
    ax.set_facecolor('#03071e')
    sns.histplot(data=df, x='processing_time_ms', hue='tool', palette=colors, kde=True, bins=15)
    plt.title('Processing Time Distribution')
    plt.xlabel('Processing Time (ms)')
    plt.ylabel('Count')
    plt.tight_layout()
    plt.savefig('temp_eval/plot_latency_hist.png', facecolor='#03071e')
    plt.close()

    # 5. Boxplot SSIM
    plt.figure(figsize=(8, 6), facecolor='#03071e')
    ax = plt.axes()
    ax.set_facecolor('#03071e')
    sns.boxplot(data=df, x='tool', y='ssim_score', hue='tool', palette=colors, legend=False)
    plt.title('SSIM Score Distribution & Variance by Tool')
    plt.ylabel('SSIM Score')
    plt.tight_layout()
    plt.savefig('temp_eval/plot_ssim_box.png', facecolor='#03071e')
    plt.close()

    # 6. Boxplot BRISQUE
    plt.figure(figsize=(8, 6), facecolor='#03071e')
    ax = plt.axes()
    ax.set_facecolor('#03071e')
    sns.boxplot(data=df, x='tool', y='brisque_score', hue='tool', palette=colors, legend=False)
    plt.title('BRISQUE Score Distribution & Variance by Tool')
    plt.ylabel('BRISQUE Score')
    plt.tight_layout()
    plt.savefig('temp_eval/plot_brisque_box.png', facecolor='#03071e')
    plt.close()

    # 7. Correlation Heatmap
    plt.figure(figsize=(8, 6), facecolor='#03071e')
    ax = plt.axes()
    ax.set_facecolor('#03071e')
    corr_df = df[['ssim_score', 'brisque_score', 'processing_time_sec']].corr()
    sns.heatmap(corr_df, annot=True, cmap='coolwarm', vmin=-1, vmax=1, fmt=".2f")
    plt.title('Correlation Heatmap')
    plt.tight_layout()
    plt.savefig('temp_eval/plot_heatmap.png', facecolor='#03071e')
    plt.close()

    # 8. User Satisfaction Pie Chart
    valid_rates = df.dropna(subset=['user_rating'])
    pos_rates = len(valid_rates[valid_rates['user_rating'] == 1])
    neg_rates = len(valid_rates[valid_rates['user_rating'] == -1])
    
    plt.figure(figsize=(6, 6), facecolor='#03071e')
    ax = plt.axes()
    ax.set_facecolor('#03071e')
    if pos_rates == 0 and neg_rates == 0:
        plt.pie([1], labels=['No Ratings'], colors=['#555555'])
    else:
        plt.pie([pos_rates, neg_rates], labels=['Thumbs Up', 'Thumbs Down'], colors=['#2a9d8f', '#e63946'], autopct='%1.1f%%', startangle=90)
    plt.title('User Satisfaction Breakdown')
    plt.tight_layout()
    plt.savefig('temp_eval/plot_satisfaction_pie.png', facecolor='#03071e')
    plt.close()


def get_case_studies(df):
    sorted_df = df.sort_values(by='ssim_score', ascending=False)
    top_3 = sorted_df.head(3).to_dict('records')
    bottom_3 = sorted_df.tail(3).to_dict('records')
    return {'top': top_3, 'bottom': bottom_3}

def download_case_study_images(case_studies):
    for i, row in enumerate(case_studies['top']):
        download_image(row['cloudinary_url'], f"temp_eval/top_{i}_orig.jpg")
        download_image(row['enhanced_url'], f"temp_eval/top_{i}_enh.jpg")
    for i, row in enumerate(case_studies['bottom']):
        download_image(row['cloudinary_url'], f"temp_eval/bottom_{i}_orig.jpg")
        download_image(row['enhanced_url'], f"temp_eval/bottom_{i}_enh.jpg")

def generate_pdf_report(df, case_studies):
    pdf_path = "reports/lumina_eval_report.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter

    # Page 1: Cover
    c.setFont("Helvetica-Bold", 32)
    c.drawString(50, height - 150, "Lumina AI")
    c.setFont("Helvetica", 20)
    c.drawString(50, height - 180, "Comprehensive Image Enhancement Evaluation Report")
    
    c.setFont("Helvetica", 14)
    c.drawString(50, height - 250, f"Date Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    c.drawString(50, height - 280, f"Total Images Evaluated: {len(df)}")
    
    avg_ssim = df['ssim_score'].mean()
    med_ssim = df['ssim_score'].median()
    avg_brisque = df['brisque_score'].mean()
    med_brisque = df['brisque_score'].median()
    avg_time = df['processing_time_ms'].mean() / 1000.0
    
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 340, "Global Summary Metrics:")
    c.setFont("Helvetica", 14)
    c.drawString(70, height - 370, f"Mean SSIM Score: {avg_ssim:.4f} (Median: {med_ssim:.4f})")
    c.drawString(70, height - 390, f"Mean BRISQUE Score: {avg_brisque:.4f} (Median: {med_brisque:.4f})")
    c.drawString(70, height - 410, f"Mean Processing Time: {avg_time:.2f} seconds")
    c.showPage()

    def draw_plot_page(title, img_path):
        c.setFont("Helvetica-Bold", 18)
        c.drawString(50, height - 50, title)
        if os.path.exists(img_path):
            c.drawImage(ImageReader(img_path), 50, height - 450, width=500, height=375)
        c.showPage()

    draw_plot_page("SSIM Scores per Tool Type", 'temp_eval/plot_ssim.png')
    draw_plot_page("SSIM Distribution & Variance (Boxplot)", 'temp_eval/plot_ssim_box.png')
    draw_plot_page("BRISQUE Scores per Tool Type", 'temp_eval/plot_brisque.png')
    draw_plot_page("BRISQUE Distribution & Variance (Boxplot)", 'temp_eval/plot_brisque_box.png')
    draw_plot_page("SSIM Score vs Processing Time", 'temp_eval/plot_scatter.png')
    draw_plot_page("Processing Latency Distribution", 'temp_eval/plot_latency_hist.png')
    draw_plot_page("Correlation Heatmap", 'temp_eval/plot_heatmap.png')
    draw_plot_page("User Satisfaction Breakdown", 'temp_eval/plot_satisfaction_pie.png')

    # Case Studies
    def draw_case_study(group, label, is_top=True):
        for i, row in enumerate(group):
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height - 50, f"{label} #{i+1} (ID: {str(row['id'])[:8]})")
            
            prefix = "top" if is_top else "bottom"
            img_orig = f"temp_eval/{prefix}_{i}_orig.jpg"
            img_enh = f"temp_eval/{prefix}_{i}_enh.jpg"
            
            if os.path.exists(img_orig):
                c.drawImage(ImageReader(img_orig), 50, height - 350, width=220, height=220, preserveAspectRatio=True)
            if os.path.exists(img_enh):
                c.drawImage(ImageReader(img_enh), 300, height - 350, width=220, height=220, preserveAspectRatio=True)
                
            c.setFont("Helvetica-Bold", 14)
            c.drawString(100, height - 380, "Original Cloudinary")
            c.drawString(350, height - 380, "Enhanced Image")
            
            c.setFont("Helvetica", 12)
            c.drawString(50, height - 420, f"Tool: {row['tool']} | SSIM: {row['ssim_score']:.4f} | BRISQUE: {row['brisque_score']:.4f}")
            c.drawString(50, height - 440, f"Processing Time: {row['processing_time_ms']}ms | User Rating: {row['user_rating']}")
            c.showPage()

    draw_case_study(case_studies['top'], "Top Performing Image (Highest SSIM)", True)
    draw_case_study(case_studies['bottom'], "Worst Performing Image (Lowest SSIM)", False)
    
    c.save()

def generate_word_report(df, case_studies):
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Lumina AI — Comprehensive Image Enhancement Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('An advanced quantitative and qualitative assessment of AI-powered image enhancement')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(f"Date: {pd.Timestamp.now().strftime('%Y-%m-%d')}\\nTotal Dataset Size: {len(df)} images")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Section 1
    doc.add_heading('Section 1 — Introduction & Methodology', level=1)
    doc.add_paragraph("Lumina AI utilizes highly advanced Generative AI and image restoration models to upscale, denoise, unblur, and relight source imagery. This report provides an in-depth statistical analysis evaluating their performance via objective structural and spatial quality metrics.")
    doc.add_paragraph("1. SSIM (Structural Similarity Index Measure): Measures the structural integrity and luminance retention comparing the enhanced output against the original source. Scores range structurally between 0 and 1.0 (perfect match).")
    doc.add_paragraph("2. BRISQUE (Blind/Referenceless Image Spatial Quality Evaluator): A no-reference IQA based on Natural Scene Statistics (NSS) computed via a pre-trained SVR model. It detects unnatural noise, distortions, banding, and synthetic ringing. A lower score signifies better perceptual naturalness.")
    doc.add_paragraph("3. Processing Latency: Recorded in milliseconds indicating the complete end-to-end trip time across the Replicate API and Cloudinary webhook callbacks.")

    # Section 2
    doc.add_heading('Section 2 — Advanced Summary Statistics', level=1)
    stats_table = doc.add_table(rows=1, cols=7)
    stats_table.style = 'Light Shading Accent 2'
    hdr = stats_table.rows[0].cells
    hdr[0].text = 'Tool'
    hdr[1].text = 'SSIM (Mean ± SD)'
    hdr[2].text = 'BRISQUE (Mean ± SD)'
    hdr[3].text = 'Min SSIM'
    hdr[4].text = 'Max SSIM'
    hdr[5].text = 'Med Time(s)'
    hdr[6].text = 'Sat. %'

    for tool in df['tool'].unique():
        subset = df[df['tool'] == tool]
        ssim_mean = subset['ssim_score'].mean()
        ssim_std = subset['ssim_score'].std() or 0
        brisque_mean = subset['brisque_score'].mean()
        brisque_std = subset['brisque_score'].std() or 0
        min_s = subset['ssim_score'].min()
        max_s = subset['ssim_score'].max()
        med_time = subset['processing_time_ms'].median() / 1000.0
        
        valid_rates = subset.dropna(subset=['user_rating'])
        pos_rates = len(valid_rates[valid_rates['user_rating'] == 1])
        total_rates = len(valid_rates[valid_rates['user_rating'].isin([1, -1])])
        sat_pct = (pos_rates / total_rates * 100) if total_rates > 0 else 0

        row = stats_table.add_row().cells
        row[0].text = str(tool).capitalize()
        row[1].text = f"{ssim_mean:.3f} ± {ssim_std:.3f}"
        row[2].text = f"{brisque_mean:.2f} ± {brisque_std:.2f}"
        row[3].text = f"{min_s:.3f}"
        row[4].text = f"{max_s:.3f}"
        row[5].text = f"{med_time:.2f}"
        row[6].text = f"{sat_pct:.1f}%"

    # Section 3
    doc.add_heading('Section 3 — Data Visualizations', level=1)
    
    doc.add_heading('3.1 Tool Variance (Boxplots)', level=2)
    doc.add_paragraph("Boxplots display the interquartile range (IQR). The box represents the middle 50% of the data, the line within the box is the median, and the extending whiskers show the statistical boundaries. Dots outside the whiskers represent outliers.")
    if os.path.exists('temp_eval/plot_ssim_box.png'):
        doc.add_picture('temp_eval/plot_ssim_box.png', width=Inches(5))
    if os.path.exists('temp_eval/plot_brisque_box.png'):
        doc.add_picture('temp_eval/plot_brisque_box.png', width=Inches(5))
        
    doc.add_heading('3.2 Correlation Analysis', level=2)
    doc.add_paragraph("The correlation heatmap visualizes relationships between parameters. A score approaching 1.0 or -1.0 implies strong correlation. For instance, testing if an increase in processing time strongly correlates to a higher quality SSIM outcome.")
    if os.path.exists('temp_eval/plot_heatmap.png'):
        doc.add_picture('temp_eval/plot_heatmap.png', width=Inches(5))

    doc.add_heading('3.3 User Satisfaction', level=2)
    doc.add_paragraph("Aggregated real-world feedback via the rating widget mechanism.")
    if os.path.exists('temp_eval/plot_satisfaction_pie.png'):
        doc.add_picture('temp_eval/plot_satisfaction_pie.png', width=Inches(5))

    # Section 4
    doc.add_heading('Section 4 — Case Studies', level=1)
    def insert_case_studies(group, prefix, is_best=True):
        label = "Best" if is_best else "Worst"
        for i, row in enumerate(group):
            p = doc.add_paragraph()
            run = p.add_run(f"Top {i+1} {label} Performing (ID: {str(row['id'])[:8]})")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0) if is_best else RGBColor(255, 0, 0)

            img_table = doc.add_table(rows=1, cols=2)
            orig_path = f"temp_eval/{prefix}_{i}_orig.jpg"
            enh_path = f"temp_eval/{prefix}_{i}_enh.jpg"
            
            if os.path.exists(orig_path):
                img_table.cell(0,0).paragraphs[0].add_run().add_picture(orig_path, width=Inches(2.5))
            if os.path.exists(enh_path):
                img_table.cell(0,1).paragraphs[0].add_run().add_picture(enh_path, width=Inches(2.5))

            doc.add_paragraph(f"Caption: Tool: {str(row['tool']).capitalize()} | SSIM: {row['ssim_score']:.4f} | BRISQUE: {row['brisque_score']:.4f} | Time: {row['processing_time_ms']}ms")

    insert_case_studies(case_studies['top'], "top", True)
    insert_case_studies(case_studies['bottom'], "bottom", False)

    # Section 5
    doc.add_heading('Section 5 — Concluding Remarks', level=1)
    ssim_mean = df['ssim_score'].mean()
    brisque_mean = df['brisque_score'].mean()
    doc.add_paragraph(f"The comprehensive benchmark across {len(df)} generations reveals a robust baseline capability. An overarching SSIM population mean of {ssim_mean:.4f} indicates excellent structural preservation, while the BRISQUE score average of {brisque_mean:.4f} tracks well against established SVR perceptual thresholds.")
    
    doc.save("reports/lumina_eval_report.docx")

if __name__ == "__main__":
    generate_reports()
